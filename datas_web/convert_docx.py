from docx import Document
import os

def docx_to_text(docx_path, output_txt_path):
    """DOCX 파일을 텍스트로 변환"""
    doc = Document(docx_path)

    with open(output_txt_path, 'w', encoding='utf-8') as f:
        for para in doc.paragraphs:
            f.write(para.text + '\n')

    print(f"텍스트 파일 저장 완료: {output_txt_path}")

def docx_to_html(docx_path, output_html_path):
    """DOCX 파일을 HTML로 변환 (간단한 버전)"""
    doc = Document(docx_path)

    html_content = ['<!DOCTYPE html>', '<html>', '<head>',
                   '<meta charset="UTF-8">',
                   '<title>TOPIK 말하기</title>',
                   '<style>',
                   'body { font-family: "Malgun Gothic", sans-serif; line-height: 1.6; margin: 20px; }',
                   'p { margin: 10px 0; }',
                   'table { border-collapse: collapse; margin: 10px 0; }',
                   '</style>',
                   '</head>', '<body>']

    for para in doc.paragraphs:
        # 스타일에 따라 다르게 처리
        if para.style.name.startswith('Heading'):
            level = para.style.name[-1] if para.style.name[-1].isdigit() else '1'
            html_content.append(f'<h{level}>{para.text}</h{level}>')
        else:
            html_content.append(f'<p>{para.text}</p>')

    # 테이블 처리
    for table in doc.tables:
        html_content.append('<table border="1">')
        for row in table.rows:
            html_content.append('<tr>')
            for cell in row.cells:
                html_content.append(f'<td>{cell.text}</td>')
            html_content.append('</tr>')
        html_content.append('</table>')

    html_content.extend(['</body>', '</html>'])

    with open(output_html_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(html_content))

    print(f"HTML 파일 저장 완료: {output_html_path}")

if __name__ == "__main__":
    docx_file = "data/TOPIK_말하기_최종수정본_y251106.docx"

    # 텍스트 변환
    docx_to_text(docx_file, "data/TOPIK_말하기_최종수정본.txt")

    # HTML 변환
    docx_to_html(docx_file, "data/TOPIK_말하기_최종수정본.html")

    print("\n변환 완료!")
